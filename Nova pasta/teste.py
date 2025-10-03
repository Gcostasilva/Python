import streamlit as st
st.set_page_config(page_title="Assistente Local", page_icon="🤖")

import os
import hashlib
import json
import pandas as pd

from langchain_community.llms import Ollama
from langchain_community.vectorstores import FAISS
from langchain.embeddings import HuggingFaceEmbeddings
from langchain.chains import ConversationalRetrievalChain
from langchain.memory import ConversationBufferMemory
from langchain.docstore.document import Document

from PyPDF2 import PdfReader
import docx

"""
Instalações:

pip install streamlit
pip install langchain langchain-community
pip install faiss-cpu
pip install sentence-transformers
pip install pypdf
pip install python-docx
pip install pandas openpyxl

Rodar:

Terminal
ollama run mistral
"""

# =============================
# Configurações globais
# =============================
VETOR_DB_PATH = "./vetor_db"
HISTORICO_PATH = "./chat_history.json"
PASTAS_ASSUNTOS = ["Estoque", "Compras", "Vendas"]

# =============================
# Funções utilitárias
# =============================
def calcular_hash(caminho_arquivo: str) -> str:
    hash_md5 = hashlib.md5()
    with open(caminho_arquivo, "rb") as f:
        for chunk in iter(lambda: f.read(4096), b""):
            hash_md5.update(chunk)
    return hash_md5.hexdigest()

# =============================
# Carregadores de documentos
# =============================
def carregar_pdfs(pasta: str):
    documentos = []
    for arquivo in os.listdir(pasta):
        if arquivo.endswith(".pdf"):
            caminho = os.path.join(pasta, arquivo)
            try:
                leitor = PdfReader(caminho)
                texto = "".join([pagina.extract_text() or "" for pagina in leitor.pages])
                documentos.append(Document(
                    page_content=texto,
                    metadata={"origem": caminho, "hash": calcular_hash(caminho)}
                ))
            except Exception as e:
                st.warning(f"Erro ao ler PDF {caminho}: {e}")
    return documentos

def carregar_word(pasta: str):
    documentos = []
    for arquivo in os.listdir(pasta):
        if arquivo.endswith((".docx", ".doc")):
            caminho = os.path.join(pasta, arquivo)
            try:
                doc = docx.Document(caminho)
                texto = "\n".join([p.text for p in doc.paragraphs if p.text.strip()])
                documentos.append(Document(
                    page_content=texto,
                    metadata={"origem": caminho, "hash": calcular_hash(caminho)}
                ))
            except Exception as e:
                st.warning(f"Erro ao ler Word {caminho}: {e}")
    return documentos

def carregar_excel(pasta: str):
    documentos = []
    for arquivo in os.listdir(pasta):
        if arquivo.endswith((".xlsx", ".xls")):
            caminho = os.path.join(pasta, arquivo)
            try:
                xls = pd.ExcelFile(caminho)
                for sheet in xls.sheet_names:
                    df = xls.parse(sheet)
                    texto = df.to_string()
                    documentos.append(Document(
                        page_content=texto,
                        metadata={"origem": f"{caminho} - {sheet}", "hash": calcular_hash(caminho)}
                    ))
            except Exception as e:
                st.warning(f"Erro ao ler Excel {caminho}: {e}")
    return documentos

def carregar_todos():
    base_docs = []
    for assunto in PASTAS_ASSUNTOS:
        pasta = f"./{assunto}"
        if os.path.exists(pasta):
            base_docs.extend(carregar_pdfs(pasta))
            base_docs.extend(carregar_word(pasta))
            base_docs.extend(carregar_excel(pasta))
    return base_docs

# =============================
# Base Vetorial (FAISS)
# =============================
def sincronizar_base():
    embeddings = HuggingFaceEmbeddings(model_name="sentence-transformers/all-MiniLM-L6-v2")

    if os.path.exists(VETOR_DB_PATH):
        vectorstore = FAISS.load_local(VETOR_DB_PATH, embeddings, allow_dangerous_deserialization=True)
        indexed_docs = {doc.metadata["origem"]: doc.metadata.get("hash", "") 
                        for doc in vectorstore.docstore._dict.values()}
    else:
        vectorstore = None
        indexed_docs = {}

    docs = carregar_todos()
    current_docs = {doc.metadata["origem"]: doc.metadata["hash"] for doc in docs}

    novos_docs = [doc for doc in docs if indexed_docs.get(doc.metadata["origem"]) != doc.metadata["hash"]]
    docs_removidos = [origem for origem in indexed_docs.keys() if origem not in current_docs]

    if novos_docs or docs_removidos or vectorstore is None:
        vectorstore = FAISS.from_documents(docs, embeddings)
        vectorstore.save_local(VETOR_DB_PATH)

    return vectorstore

# =============================
# Histórico persistente
# =============================
def carregar_historico():
    if os.path.exists(HISTORICO_PATH):
        with open(HISTORICO_PATH, "r", encoding="utf-8") as f:
            return json.load(f)
    return []

def salvar_historico(chat_history):
    with open(HISTORICO_PATH, "w", encoding="utf-8") as f:
        json.dump(chat_history, f, indent=2, ensure_ascii=False)

def inicializar_memoria():
    memoria = ConversationBufferMemory(memory_key="chat_history", return_messages=True)
    historico = carregar_historico()
    for msg in historico:
        if msg["role"] == "user":
            memoria.chat_memory.add_user_message(msg["content"])
        else:
            memoria.chat_memory.add_ai_message(msg["content"])
    return memoria

# =============================
# Interface Streamlit
# =============================
st.title("🤖 Assistente de Documentos Locais")
st.write("Base: PDFs, Word e Excel em pastas **Estoque, Compras, Vendas**")

# Atualizar base manualmente
if st.button("🔄 Atualizar base de conhecimento"):
    with st.spinner("Sincronizando documentos..."):
        vectorstore = sincronizar_base()
    st.success("Base atualizada com sucesso!")

# Inicialização
vectorstore = sincronizar_base()
retriever = vectorstore.as_retriever(search_kwargs={"k": 3})

if "memory" not in st.session_state:
    st.session_state.memory = inicializar_memoria()

llm = Ollama(model="mistral")
qa_chain = ConversationalRetrievalChain.from_llm(
    llm=llm,
    retriever=retriever,
    memory=None,
    return_source_documents=True  # NÃO PASSAR memória aqui
)

# Caixa de chat
st.subheader("💬 Chat com a IA")
query = st.chat_input("Digite sua pergunta...")

if query:
    with st.spinner("Consultando IA..."):
        result = qa_chain({"question": query})

        # Salvar manualmente na memória
        st.session_state.memory.chat_memory.add_user_message(query)
        st.session_state.memory.chat_memory.add_ai_message(result["answer"])

    # Exibir mensagens
    st.chat_message("user").write(query)
    st.chat_message("assistant").write(result["answer"])

    # Atualizar histórico persistente
    historico = carregar_historico()
    historico.append({"role": "user", "content": query})
    historico.append({"role": "assistant", "content": result["answer"]})
    salvar_historico(historico)

    # Fontes consultadas
    with st.expander("📄 Fontes consultadas"):
        for doc in result["source_documents"]:
            st.write("-", doc.metadata["origem"])

# Exibir histórico completo
if st.checkbox("📜 Mostrar histórico de conversas"):
    historico = carregar_historico()
    for msg in historico:
        role = "👤 Você" if msg["role"] == "user" else "🤖 Assistente"
        st.write(f"**{role}:** {msg['content']}")
